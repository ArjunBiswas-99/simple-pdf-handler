"""
OCR Coordinator - Orchestrates the complete OCR workflow.

Combines all OCR components (engine, processor, detector, analyzer) and provides
background processing with progress tracking. This is the main entry point for
all OCR operations.
"""

import time
from typing import List, Optional, Dict, Any, Callable
from pathlib import Path
from PySide6.QtCore import QThread, Signal
import fitz

from .ocr_engine import OCREngine, OCRResult
from .image_processor import ImageProcessor
from .text_extractor import TextExtractor, PageOCRResult
from .table_detector import TableDetector, TableStructure
from .pdf_text_layer import PDFTextLayer


class ScanDetector:
    """Detects if a PDF is a scanned document without text layer."""
    
    @staticmethod
    def is_scanned_pdf(pdf_path: str, sample_pages: int = 5) -> bool:
        """
        Check if PDF appears to be scanned (no text layer).
        
        Uses multiple heuristics to determine if a PDF is scanned:
        - Text content analysis (amount of extractable text)
        - Image to text ratio
        - Page composition (images vs text objects)
        
        Args:
            pdf_path: Path to PDF file
            sample_pages: Number of pages to sample (min 5 for accuracy)
            
        Returns:
            True if document appears to be scanned (image-only)
        """
        try:
            doc = fitz.open(pdf_path)
            num_pages = len(doc)
            pages_to_check = min(sample_pages, num_pages)
            
            scanned_indicators = 0
            total_checks = 0
            
            print(f"DEBUG ScanDetector: Analyzing {pages_to_check} pages from {num_pages} total")
            
            for i in range(pages_to_check):
                # Sample pages evenly distributed throughout document
                page_num = int(i * num_pages / pages_to_check) if num_pages > pages_to_check else i
                page = doc[page_num]
                
                # Check 1: Text content
                text = page.get_text().strip()
                text_length = len(text)
                
                # Check 2: Image presence
                images = page.get_images()
                has_large_images = len(images) > 0
                
                # Check 3: Text blocks (more reliable than raw text)
                text_blocks = page.get_text("blocks")
                meaningful_blocks = [b for b in text_blocks if len(b[4].strip()) > 10]
                
                print(f"DEBUG: Page {page_num} - Text: {text_length} chars, Images: {len(images)}, Text blocks: {len(meaningful_blocks)}")
                
                # Scoring: A scanned page typically has:
                # - Very little extractable text (< 100 chars)
                # - At least one large image covering most of the page
                # - Few or no text blocks
                
                page_score = 0
                
                if text_length < 100:
                    page_score += 1
                    print(f"DEBUG:   → Little text (scanned indicator)")
                
                if has_large_images and len(meaningful_blocks) < 2:
                    page_score += 1
                    print(f"DEBUG:   → Has images but few text blocks (scanned indicator)")
                
                if text_length < 20 and has_large_images:
                    page_score += 1  # Strong indicator
                    print(f"DEBUG:   → Almost no text + images (strong scanned indicator)")
                
                # If page has substantial text (>200 chars) and text blocks, it's definitely not scanned
                if text_length > 200 and len(meaningful_blocks) > 0:
                    page_score = 0  # Override - clearly has text layer
                    print(f"DEBUG:   → Substantial text found (NOT scanned)")
                
                if page_score >= 2:
                    scanned_indicators += 1
                
                total_checks += 1
            
            doc.close()
            
            # Consider scanned if majority of sampled pages appear scanned
            is_scanned = scanned_indicators > (total_checks * 0.6)
            
            print(f"DEBUG ScanDetector: Result = {is_scanned} ({scanned_indicators}/{total_checks} pages appear scanned)")
            
            return is_scanned
            
        except Exception as e:
            print(f"DEBUG ScanDetector ERROR: {e}")
            return False


class ConfidenceAnalyzer:
    """Analyzes OCR confidence and identifies suspicious words."""
    
    def __init__(self, threshold: float = 0.75):
        """
        Initialize analyzer.
        
        Args:
            threshold: Confidence threshold for suspicious words
        """
        self.threshold = threshold
    
    def analyze_results(
        self,
        ocr_results: List[PageOCRResult]
    ) -> Dict[str, Any]:
        """
        Analyze OCR results for confidence and accuracy.
        
        Args:
            ocr_results: List of page OCR results
            
        Returns:
            Dictionary with analysis statistics
        """
        total_words = 0
        total_confidence = 0.0
        suspicious_words = []
        
        for page_result in ocr_results:
            for text_block in page_result.text_blocks:
                words = text_block.text.split()
                total_words += len(words)
                total_confidence += text_block.confidence
                
                if text_block.confidence < self.threshold:
                    suspicious_words.append({
                        'text': text_block.text,
                        'confidence': text_block.confidence,
                        'page': page_result.page_number,
                        'bbox': text_block.bbox
                    })
        
        # Calculate average confidence safely
        all_blocks = [b for r in ocr_results for b in r.text_blocks]
        avg_confidence = (total_confidence / len(all_blocks)) if all_blocks else 0.0
        
        return {
            'total_words': total_words,
            'avg_confidence': avg_confidence,
            'suspicious_count': len(suspicious_words),
            'suspicious_words': suspicious_words,
            'pages_processed': len(ocr_results)
        }
    
    def get_suggestions(self, word: str, context: str) -> List[str]:
        """
        Get correction suggestions for a word (simple implementation).
        
        Args:
            word: Word to correct
            context: Surrounding context
            
        Returns:
            List of suggested corrections
        """
        # Simple placeholder - could integrate spell checker
        return [word, word.lower(), word.upper(), word.capitalize()]


class ExportHandler:
    """Handles exporting OCR results to various formats."""
    
    @staticmethod
    def export_to_text(
        ocr_results: List[PageOCRResult],
        output_path: str
    ) -> bool:
        """Export to plain text file."""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                for result in ocr_results:
                    f.write(f"=== Page {result.page_number + 1} ===\n\n")
                    f.write(result.full_text)
                    f.write("\n\n")
            return True
        except:
            return False
    
    @staticmethod
    def export_to_word(
        ocr_results: List[PageOCRResult],
        output_path: str
    ) -> bool:
        """Export to Word document."""
        try:
            from docx import Document
            
            doc = Document()
            for result in ocr_results:
                doc.add_heading(f'Page {result.page_number + 1}', level=2)
                doc.add_paragraph(result.full_text)
                doc.add_page_break()
            
            doc.save(output_path)
            return True
        except:
            return False
    
    @staticmethod
    def export_tables_to_excel(
        tables: List[TableStructure],
        output_path: str
    ) -> bool:
        """Export detected tables to Excel."""
        try:
            from openpyxl import Workbook
            
            wb = Workbook()
            ws = wb.active
            ws.title = "OCR Tables"
            
            current_row = 1
            for i, table in enumerate(tables):
                ws.cell(current_row, 1, f"Table {i+1}")
                current_row += 2
                
                for cell in table.cells:
                    ws.cell(current_row + cell.row, cell.col + 1, cell.text)
                
                current_row += table.rows + 2
            
            wb.save(output_path)
            return True
        except:
            return False


class OCRWorker(QThread):
    """Background worker thread for OCR processing."""
    
    # Signals
    progress_updated = Signal(int, int, str)  # current, total, message
    page_completed = Signal(int, dict)  # page_num, statistics
    ocr_completed = Signal(list, dict)  # results, statistics
    error_occurred = Signal(str)  # error message
    
    def __init__(
        self,
        pdf_path: str,
        language: str = 'en',
        page_range: Optional[tuple] = None,
        apply_preprocessing: bool = True
    ):
        """
        Initialize OCR worker.
        
        Args:
            pdf_path: Path to PDF file
            language: OCR language code
            page_range: Optional (start, end) page numbers
            apply_preprocessing: Whether to apply image enhancement
        """
        super().__init__()
        self.pdf_path = pdf_path
        self.language = language
        self.page_range = page_range
        self.apply_preprocessing = apply_preprocessing
        self._cancelled = False
    
    def run(self):
        """Execute OCR processing in background thread."""
        try:
            # Initialize components
            self.progress_updated.emit(0, 100, "Initializing OCR engine...")
            print(f"DEBUG OCRWorker: Starting OCR on {self.pdf_path}")
            print(f"DEBUG: Language: {self.language}")
            
            ocr_engine = OCREngine(self.language)
            image_processor = ImageProcessor()
            text_extractor = TextExtractor(ocr_engine, image_processor)
            
            # Open PDF
            doc = fitz.open(self.pdf_path)
            print(f"DEBUG: Opened PDF - {len(doc)} pages")
            
            # Determine page range
            if self.page_range:
                start, end = self.page_range
            else:
                start, end = 0, len(doc)
            
            total_pages = end - start
            results = []
            
            print(f"DEBUG: Processing pages {start} to {end} ({total_pages} pages)")
            
            # Process each page
            for i, page_num in enumerate(range(start, end)):
                if self._cancelled:
                    print("DEBUG: OCR cancelled by user")
                    break
                
                # Update progress
                progress = int((i / total_pages) * 100)
                self.progress_updated.emit(
                    i + 1,
                    total_pages,
                    f"Processing page {i + 1} of {total_pages}..."
                )
                
                # Extract text from page
                result = text_extractor.extract_text_from_page(
                    doc,
                    page_num,
                    preprocess=self.apply_preprocessing
                )
                
                # Debug output
                print(f"DEBUG: Page {page_num} OCR complete - {len(result.text_blocks)} text blocks, {result.word_count} words")
                if result.text_blocks:
                    print(f"DEBUG: Sample text: {result.text_blocks[0].text[:50]}...")
                
                results.append(result)
                
                # Emit page completion
                stats = text_extractor.get_page_statistics(result)
                self.page_completed.emit(page_num, stats)
            
            doc.close()
            
            if not self._cancelled:
                # Analyze overall results
                print(f"DEBUG: OCR complete, analyzing {len(results)} page results")
                analyzer = ConfidenceAnalyzer()
                analysis = analyzer.analyze_results(results)
                
                print(f"DEBUG: Analysis - Total words: {analysis['total_words']}, Avg confidence: {analysis['avg_confidence']:.2f}")
                
                # Emit completion
                self.ocr_completed.emit(results, analysis)
            
        except Exception as e:
            error_msg = str(e)
            print(f"DEBUG OCR ERROR: {error_msg}")
            import traceback
            traceback.print_exc()
            self.error_occurred.emit(error_msg)
    
    def cancel(self):
        """Cancel OCR processing."""
        self._cancelled = True


class OCRCoordinator:
    """
    Main coordinator for all OCR operations.
    
    Provides high-level interface to OCR functionality with
    progress tracking and error handling.
    """
    
    def __init__(self):
        """Initialize OCR coordinator with all components."""
        self.ocr_engine = None
        self.image_processor = ImageProcessor()
        self.text_extractor = None
        self.table_detector = TableDetector()
        self.pdf_text_layer = PDFTextLayer()
        self.scan_detector = ScanDetector()
        self.confidence_analyzer = ConfidenceAnalyzer()
        self.export_handler = ExportHandler()
    
    def initialize_engine(self, language: str = 'en'):
        """Initialize OCR engine with specified language."""
        self.ocr_engine = OCREngine(language)
        self.text_extractor = TextExtractor(self.ocr_engine, self.image_processor)
    
    def process_document(
        self,
        pdf_path: str,
        output_path: str,
        language: str = 'en',
        page_range: Optional[tuple] = None,
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """
        Process entire document with OCR.
        
        Args:
            pdf_path: Input PDF path
            output_path: Output PDF path
            language: OCR language
            page_range: Optional page range
            progress_callback: Progress callback function
            
        Returns:
            Dictionary with results and statistics
        """
        if self.ocr_engine is None or self.ocr_engine.language != language:
            self.initialize_engine(language)
        
        # Extract text from all pages
        ocr_results = self.text_extractor.extract_text_from_document(
            pdf_path,
            page_range,
            preprocess=True,
            progress_callback=progress_callback
        )
        
        # Create searchable PDF
        success = self.pdf_text_layer.create_searchable_pdf(
            pdf_path,
            output_path,
            ocr_results,
            compress=True
        )
        
        # Analyze results
        analysis = self.confidence_analyzer.analyze_results(ocr_results)
        
        return {
            'success': success,
            'output_path': output_path if success else None,
            'statistics': analysis
        }
    
    def is_scanned_document(self, pdf_path: str) -> bool:
        """Check if document is scanned."""
        return self.scan_detector.is_scanned_pdf(pdf_path)
    
    def get_document_info(self, pdf_path: str) -> Dict[str, Any]:
        """Get document information."""
        if self.text_extractor is None:
            self.initialize_engine()
        return self.text_extractor.get_document_info(pdf_path)
